import whisper
from docx import Document
from fpdf import FPDF
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from pyAudioAnalysis import audioSegmentation

# Load model
model = whisper.load_model("base")

# Load audio
audio = whisper.load_audio("audio.mp3")
audio = whisper.pad_or_trim(audio) ## trims it to 30 seconds

# Make log-Mel spectrogram and move to the same device as the model
mel = whisper.log_mel_spectrogram(audio).to(model.device)

# Detect the spoken language
_, probs = model.detect_language(mel)
print(f"Detected language: {max(probs, key=probs.get)}")

# Perform speaker diarization
def perform_speaker_diarization(audio):
    # Convert audio to mono and resample to 8 kHz
    mono_audio = whisper.to_mono(audio)
    resampled_audio = whisper.resample(mono_audio, 8000)

    # Perform speaker diarization using pyAudioAnalysis
    segments = audioSegmentation.speaker_diarization(resampled_audio)

    # Retrieve the transcriptions and speaker information
    transcriptions = []
    for segment in segments:
        start_time, end_time, speaker_label = segment
        transcript = whisper.decode(model, mel[:, int(start_time * 100):int(end_time * 100)]).text
        transcriptions.append((speaker_label, transcript))

    return transcriptions

# Call the function to perform speaker diarization
transcriptions = perform_speaker_diarization(audio) ## replace audio with the audio name file

document = Document()
for speaker_label, transcript in transcriptions:
    document.add_paragraph(f"Speaker {speaker_label}: {transcript}")

# Decode the audio
options = whisper.DecodingOptions()
result = whisper.decode(model, mel, options)

# Print the recognized text
print(result.text)

# Generate transcription file
first_name = input("Enter participant's first name: ")
last_name = input("Enter participant's last name: ")
week_number = int(input("Enter participant's interview week number: "))

transcription_name_word = f"{first_name}{last_name} Week {week_number}.docx"
transcription_name_pdf = f"{first_name}{last_name} Week {week_number}.pdf"

# Save the recognized text to as a Word document

document = Document()
document.save(transcription_name_word)
print(f"Transcription saved as {transcription_name_word} Word document.")

# Save the recognized text to as a PDF document

pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.cell(0, 10, result.text)
pdf.output(pdf_name)
print(f"Transcription saved as {transcription_name_pdf} PDF document.")

# Upload transcription to Google Drive

gauth = GoogleAuth()
drive = GoogleDrive(gauth)
file1 = drive.CreateFile({'title': transcription_name_word})
file1.SetContentFile(transcription_name_word)
file1.Upload()
print(f"Transcription saved to Google Drive as {transcription_name_word}.")

file2 = drive.CreateFile({'title': transcription_name_pdf})
file2.SetContentFile(transcription_name_pdf)
file2.Upload()
print(f"Transcription saved to Google Drive as {transcription_name_pdf}.")
